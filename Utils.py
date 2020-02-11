import datetime, time
import json
import re
import sqlite3
import urllib.request
from urllib.parse import urljoin
from docx import Document
from docx.shared import Cm
import os


"""记录创建MySql表用的语句
create table file_articles (id int not null primary key auto_increment,
title varchar(300) not null,
url varchar(100),
keywords varchar(100),
author varchar(500),
published_time varchar(50),
picture_url_list varchar(500),
content mediumtext
);
"""

def date_limit(date, days_delta=-2, date_set=""):
    '''
    :param date: 一个datetime字符串，格式%Y-%m-%d %H:%M:%S
    :param days_delta:天数
    :param date_set:手动设定比较值
    :return:bool值，是否是date_set之后的一个时间
    '''
    if datetime.datetime.strptime(date, "%Y-%m-%d %H:%M:%S") > datetime.datetime.now() + datetime.timedelta(days=days_delta):
        return True
    else:
        return False


class Exportor(object):
    yesterday = datetime.date.today() + datetime.timedelta(days=-1)
    data_limit = datetime.datetime.strptime(str(yesterday) + " 22:30:00", "%Y-%m-%d %H:%M:%S")

    def __init__(self):
        # 连接SQLite
        self.connect = sqlite3.connect('./data/WeChat_ThePublic.db')
        print("Opened database successfully")
        self.cursor = self.connect.cursor()

    def get_article(self, data_limit="all"):
        """
        :param data_limit:可以设置为today yesterday all
        :return:
        """
        c = self.cursor.execute("""select count(*) from file_articles where datetime(published_time) > datetime("%s")""" % self.data_limit)
        print("符合时间限制的新闻有:", c.fetchall()[0][0])
        c = self.cursor.execute("""select count(*) from file_articles where datetime(published_time) > datetime("%s") and length(content) < 1000""" % self.data_limit)
        print("符合时间+字数限制的新闻有:", c.fetchall()[0][0])
        curosr = self.cursor.execute("""select * from file_articles where datetime(published_time) > datetime("%s") order by length(content) asc""" % self.data_limit)
        for row in curosr:
            ID = row[0]
            title = row[1]
            url = row[2]
            keywords = row[3]
            author = row[4]
            published_time = row[5]
            picture_url_list = json.loads(row[6])
            content = row[7]

            img = None
            picture_list = []
            if picture_url_list:
                if len(picture_url_list) > 1:
                    url = picture_url_list[1]
                else:
                    url = picture_url_list[0]
                response = urllib.request.urlopen(urljoin("http://", url))
                img = response.read()
                picture_list.append(img)
            yield (title, published_time, url, content, img)

    def export_docx(self, file_path=""):
        if file_path:
            document = Document(file_path)
        else:
            document = Document()
        for i, item in enumerate(self.get_article()):
            for i in range(len(item) - 1):
                pa = document.add_paragraph(item[i])
            if item[-1] != None:
                with open("./data/temp.jpg", "wb") as fo:
                    fo.write(item[-1])
                document.add_picture("./data/temp.jpg", width=Cm(5))
            document.add_page_break()
        document.save("./data/WeChat_ThePublic.docx")


if __name__ == '__main__':
    ex = Exportor()
    ex.export_docx()

    timeArray = time.localtime(1580173043)  # 1970秒数
    otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
    dt1 = datetime.datetime.strptime(otherStyleTime, "%Y-%m-%d %H:%M:%S")

